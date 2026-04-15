"""
Cyberwijzer Rapport Generator
=============================
Zet de JSON-output van agent.py om naar een professioneel
Word-document (.docx) in de stijl van het Cyberwijzer-template.

Gebruik:
    python rapport.py risicoanalyse_versiebeheer_met_git.json
    python rapport.py rapport.json --output mijn_rapport.docx

Vereisten:
    pip install python-docx
"""

import argparse
import json
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ---------------------------------------------------------------------------
# Kleuren en stijlen
# ---------------------------------------------------------------------------

BLAUW = RGBColor(0x2E, 0x75, 0xB6)
DONKERGRIJS = RGBColor(0x33, 0x33, 0x33)
LICHTGRIJS_HEX = "F2F2F2"
BLAUW_HEX = "D5E8F0"
HEADER_HEX = "2E75B6"


# ---------------------------------------------------------------------------
# Helper functies
# ---------------------------------------------------------------------------

def set_cell_shading(cell, hex_color):
    """Stel de achtergrondkleur van een tabelcel in."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_paragraph(doc, text, style="Normal", bold=False, size=None, color=None, space_after=None):
    """Voeg een gestileerde paragraaf toe."""
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    if bold:
        run.bold = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet_list(doc, items, indent_level=0):
    """Voeg een lijst met opsommingstekens toe."""
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        if indent_level > 0:
            p.paragraph_format.left_indent = Cm(1.27 * (indent_level + 1))


def add_table_with_header(doc, headers, rows, col_widths=None):
    """Maak een gestileerde tabel met header-rij."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header-rij
    header_row = table.rows[0]
    for i, header_text in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header_text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, HEADER_HEX)

    # Data-rijen
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, cell_text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = str(cell_text)
            cell.paragraphs[0].runs[0].font.size = Pt(10) if cell.paragraphs[0].runs else None
            # Afwisselende rijkleuren
            if row_idx % 2 == 1:
                set_cell_shading(cell, LICHTGRIJS_HEX)

    # Kolombreedtes
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table


# ---------------------------------------------------------------------------
# Rapport secties
# ---------------------------------------------------------------------------

def add_cover_page(doc, data):
    """Voeg een voorpagina toe."""
    meta = data["metadata"]

    # Lege ruimte
    for _ in range(6):
        doc.add_paragraph()

    # Project naam
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(meta.get("project", "CYBERWIJZER"))
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = BLAUW

    # Titel
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(meta.get("titel", "Risicoanalyse"))
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = DONKERGRIJS

    # Thema
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(meta.get("thema", ""))
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = BLAUW

    # Metadata tabel
    doc.add_paragraph()
    doc.add_paragraph()

    info_items = [
        ("Auteur", meta.get("auteur", "")),
        ("Datum", meta.get("datum", "")),
        ("Templateversie", meta.get("template_versie", "")),
        ("Versie", meta.get("versie", "")),
    ]

    table = doc.add_table(rows=len(info_items), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(info_items):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[0].width = Cm(5)
        row.cells[1].text = value
        row.cells[1].width = Cm(8)

    doc.add_page_break()


def add_versiebeheer(doc, data):
    """Voeg de versiebeheertabel toe."""
    doc.add_heading("Versiebeheer", level=1)

    versies = data.get("versiebeheer", [])
    if versies:
        headers = ["Versie", "Datum", "Wijzigingen", "Auteur"]
        rows = [
            [v.get("versie", ""), v.get("datum", ""),
             v.get("wijzigingen", ""), v.get("auteur", "")]
            for v in versies
        ]
        add_table_with_header(doc, headers, rows, col_widths=[2, 3, 7, 4])

    doc.add_page_break()


def add_beschrijving_thema(doc, data):
    """Sectie 1: Beschrijving van het thema."""
    thema = data.get("beschrijving_thema", {})

    doc.add_heading("Beschrijving van het thema", level=1)
    doc.add_paragraph(thema.get("inleiding", ""))

    gebruik = thema.get("gebruik", {})
    if gebruik.get("toepassingen"):
        doc.add_heading("Gebruik in softwareontwikkeling", level=2)
        add_bullet_list(doc, gebruik["toepassingen"])

    if gebruik.get("rol_in_processen"):
        add_styled_paragraph(doc, "Rol in processen:", bold=True, space_after=4)
        add_bullet_list(doc, gebruik["rol_in_processen"])

    if thema.get("conclusie"):
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run("Conclusie: ")
        run.bold = True
        p.add_run(thema["conclusie"])


def add_onderzoek(doc, data):
    """Sectie 2: Onderzoek beschikbare informatie."""
    doc.add_heading("Onderzoek beschikbare informatie", level=1)

    for item in data.get("onderzoek", []):
        doc.add_heading(item.get("titel", ""), level=2)
        doc.add_paragraph(item.get("bevindingen", ""))

        if item.get("bewezen"):
            p = doc.add_paragraph()
            run = p.add_run("Bewezen: ")
            run.bold = True
            p.add_run(item["bewezen"])

        if item.get("afgeleide_conclusie"):
            p = doc.add_paragraph()
            run = p.add_run("Afgeleide conclusie: ")
            run.bold = True
            p.add_run(item["afgeleide_conclusie"])

        if item.get("bronnen"):
            p = doc.add_paragraph()
            run = p.add_run("Bronnen: ")
            run.italic = True
            p.add_run(", ".join(item["bronnen"]))


def add_consequenties(doc, data):
    """Sectie 3: Consequenties en échte voorbeelden."""
    doc.add_heading("Consequenties en échte voorbeelden", level=1)

    for item in data.get("consequenties_en_voorbeelden", []):
        doc.add_heading(item.get("titel", ""), level=2)

        if item.get("onderzoek"):
            p = doc.add_paragraph()
            run = p.add_run("Onderzoek: ")
            run.bold = True
            p.add_run(item["onderzoek"])

        if item.get("wat_ging_mis"):
            add_styled_paragraph(doc, "Wat ging mis:", bold=True, space_after=2)
            add_bullet_list(doc, item["wat_ging_mis"])

        if item.get("impact"):
            add_styled_paragraph(doc, "Impact:", bold=True, space_after=2)
            add_bullet_list(doc, item["impact"])

    if data.get("consequenties_conclusie"):
        doc.add_heading("Conclusie", level=2)
        doc.add_paragraph(data["consequenties_conclusie"])


def add_risicos(doc, data):
    """Sectie 4: Risico's met matrix en details."""
    risicos = data.get("risicos", {})
    lijst = risicos.get("lijst", [])

    doc.add_heading("Risico's", level=1)

    # Risicomatrix tabel
    if lijst:
        headers = ["Nr", "Risico", "Kans", "Impact", "Score", "Ranking"]
        rows = [
            [
                str(r.get("nr", "")),
                r.get("titel", ""),
                str(r.get("kans", "")),
                str(r.get("impact", "")),
                str(r.get("score", "")),
                str(r.get("ranking", "")),
            ]
            for r in lijst
        ]
        add_table_with_header(doc, headers, rows, col_widths=[1.5, 6, 2, 2, 2, 2])

    # Schaaluitleg
    doc.add_paragraph()
    matrix = risicos.get("risico_matrix", {})
    kans = matrix.get("kans_schaal", {})
    impact = matrix.get("impact_schaal", {})

    if kans or impact:
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"

        # Kans kolom
        cell_kans = tbl.rows[0].cells[0]
        cell_kans.text = ""
        p = cell_kans.paragraphs[0]
        run = p.add_run("Kans")
        run.bold = True
        for k, v in kans.items():
            cell_kans.add_paragraph(f"{v}", style="List Bullet")

        # Impact kolom
        cell_impact = tbl.rows[0].cells[1]
        cell_impact.text = ""
        p = cell_impact.paragraphs[0]
        run = p.add_run("Impact")
        run.bold = True
        for k, v in impact.items():
            cell_impact.add_paragraph(f"{v}", style="List Bullet")

    # Detailbeschrijvingen per risico
    for r in lijst:
        doc.add_heading(f"Risico {r.get('nr', '')}: {r.get('titel', '')}", level=2)

        if r.get("oorzaak"):
            add_styled_paragraph(doc, "Oorzaak", bold=True, space_after=2)
            add_bullet_list(doc, r["oorzaak"])

        if r.get("gevolg"):
            add_styled_paragraph(doc, "Gevolg", bold=True, space_after=2)
            add_bullet_list(doc, r["gevolg"])

        if r.get("impact_detail"):
            add_styled_paragraph(doc, "Impact", bold=True, space_after=2)
            add_bullet_list(doc, r["impact_detail"])

        if r.get("opmerking"):
            p = doc.add_paragraph()
            run = p.add_run(r["opmerking"])
            run.italic = True


def add_maatregelen(doc, data):
    """Sectie 5: Maatregelen."""
    maatregelen = data.get("maatregelen", {})

    doc.add_heading("Maatregelen", level=1)

    for categorie, label in [
        ("technisch", "Technische maatregelen"),
        ("proces", "Procesmaatregelen"),
        ("organisatorisch", "Organisatorische maatregelen"),
    ]:
        cat = maatregelen.get(categorie, {})
        if cat:
            doc.add_heading(label, level=2)
            if cat.get("items"):
                add_bullet_list(doc, cat["items"])
            if cat.get("onderbouwing"):
                p = doc.add_paragraph()
                run = p.add_run(cat["onderbouwing"])
                run.italic = True


def add_basisvaardigheden(doc, data):
    """Sectie 6: Security basisvaardigheden."""
    bv = data.get("basisvaardigheden", {})

    doc.add_heading("Security basisvaardigheden (digitale hygiëne)", level=1)

    if bv.get("inleiding"):
        doc.add_paragraph(bv["inleiding"])

    # Bewustzijn
    bewustzijn = bv.get("bewustzijn", {})
    if bewustzijn:
        doc.add_heading("Bewustzijn (Security Awareness)", level=2)

        if bewustzijn.get("doel"):
            doc.add_heading("Doel", level=3)
            doc.add_paragraph(bewustzijn["doel"])

        if bewustzijn.get("minimale_kennis"):
            doc.add_heading("Minimale kennis en inzicht", level=3)
            doc.add_paragraph("Studenten moeten kunnen uitleggen:")
            add_bullet_list(doc, bewustzijn["minimale_kennis"])

        if bewustzijn.get("herkenningsvaardigheid"):
            doc.add_heading("Herkenningsvaardigheid", level=3)
            doc.add_paragraph("Student kan risico's herkennen in situaties zoals:")
            add_bullet_list(doc, bewustzijn["herkenningsvaardigheid"])

        if bewustzijn.get("toetsbare_leeruitkomsten"):
            doc.add_heading("Toetsbare leeruitkomsten", level=3)
            add_bullet_list(doc, bewustzijn["toetsbare_leeruitkomsten"])

    # Vaardigheden
    vaardigheden = bv.get("vaardigheden", {})
    if vaardigheden:
        doc.add_heading("Vaardigheden (Technische en ontwerpvaardigheden)", level=2)

        if vaardigheden.get("doel"):
            doc.add_heading("Doel", level=3)
            doc.add_paragraph(vaardigheden["doel"])

        for cat in vaardigheden.get("categorieen", []):
            doc.add_heading(cat.get("titel", ""), level=3)

            if cat.get("student_kan"):
                doc.add_paragraph("Student kan:")
                add_bullet_list(doc, cat["student_kan"])

            if cat.get("toetsbaar"):
                add_styled_paragraph(doc, "Toetsbaar:", bold=True, space_after=2)
                add_bullet_list(doc, cat["toetsbaar"])

    # Gedrag
    gedrag = bv.get("gedrag", {})
    if gedrag:
        doc.add_heading("Gedrag (Security mindset en professioneel handelen)", level=2)

        if gedrag.get("doel"):
            doc.add_heading("Doel", level=3)
            doc.add_paragraph(gedrag["doel"])

        if gedrag.get("gewenst_gedrag"):
            doc.add_heading("Gewenst gedrag", level=3)
            add_bullet_list(doc, gedrag["gewenst_gedrag"])

        if gedrag.get("observeerbaar_gedrag"):
            doc.add_heading("Concreet observeerbaar gedrag", level=3)
            add_bullet_list(doc, gedrag["observeerbaar_gedrag"])

        if gedrag.get("toetsbare_leeruitkomsten"):
            doc.add_heading("Toetsbare leeruitkomsten", level=3)
            add_bullet_list(doc, gedrag["toetsbare_leeruitkomsten"])

    # Minimale competentie-eisen
    if bv.get("minimale_competentie_eisen"):
        doc.add_heading("Minimale competentie-eisen (ondergrens voor beoordeling)", level=2)
        doc.add_paragraph("Een student moet minimaal:")
        for i, eis in enumerate(bv["minimale_competentie_eisen"], 1):
            doc.add_paragraph(f"{i}. {eis}")


def add_samenvatting(doc, data):
    """Sectie 7: Samenvatting."""
    samenvatting = data.get("samenvatting", {})

    doc.add_heading("Samenvatting", level=1)

    if samenvatting.get("belangrijkste_risicos"):
        add_styled_paragraph(doc, "Belangrijkste risico's", bold=True, space_after=2)
        add_bullet_list(doc, samenvatting["belangrijkste_risicos"])

    if samenvatting.get("belangrijkste_maatregelen"):
        add_styled_paragraph(doc, "Belangrijkste maatregelen", bold=True, space_after=2)
        add_bullet_list(doc, samenvatting["belangrijkste_maatregelen"])

    if samenvatting.get("kerninzicht"):
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run("Kerninzicht: ")
        run.bold = True
        p.add_run(samenvatting["kerninzicht"])


def add_bronnen(doc, data):
    """Sectie 8: Bronnen."""
    doc.add_heading("Bronnen", level=1)

    for bron in data.get("bronnen", []):
        ref = bron.get("referentie", "")
        url = bron.get("url")

        if url:
            doc.add_paragraph(f"{ref} Geraadpleegd via {url}")
        else:
            doc.add_paragraph(ref)


# ---------------------------------------------------------------------------
# Documentstijlen configureren
# ---------------------------------------------------------------------------

def configure_styles(doc):
    """Stel de standaard documentstijlen in."""
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = DONKERGRIJS
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # Heading stijlen
    for level, size in [(1, 16), (2, 13), (3, 11)]:
        heading_style = doc.styles[f"Heading {level}"]
        heading_style.font.name = "Calibri"
        heading_style.font.size = Pt(size)
        heading_style.font.color.rgb = BLAUW
        heading_style.font.bold = True
        heading_style.paragraph_format.space_before = Pt(12)
        heading_style.paragraph_format.space_after = Pt(6)

    # Paginamarges
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


# ---------------------------------------------------------------------------
# Hoofdfunctie
# ---------------------------------------------------------------------------

def generate_report(json_path: str, output_path: str | None = None) -> str:
    """Genereer een Word-rapport vanuit een JSON-bestand."""

    # Lees JSON
    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    # Bepaal output pad
    if not output_path:
        thema = data.get("metadata", {}).get("thema", "rapport")
        safe_name = thema.lower().replace(" ", "_")
        safe_name = "".join(c for c in safe_name if c.isalnum() or c == "_")
        output_path = f"Cyberwijzer_Risicoanalyse_{safe_name}.docx"

    # Maak document
    doc = Document()
    configure_styles(doc)

    # Bouw het rapport op
    add_cover_page(doc, data)
    add_versiebeheer(doc, data)
    add_beschrijving_thema(doc, data)
    add_onderzoek(doc, data)
    add_consequenties(doc, data)
    add_risicos(doc, data)
    add_maatregelen(doc, data)
    add_basisvaardigheden(doc, data)
    add_samenvatting(doc, data)
    add_bronnen(doc, data)

    # Sla op
    doc.save(output_path)
    return output_path


def main():
    parser = argparse.ArgumentParser(
        description="Cyberwijzer Rapport Generator — "
        "zet risicoanalyse-JSON om naar een Word-document.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Voorbeelden:
  python rapport.py risicoanalyse_versiebeheer_met_git.json
  python rapport.py rapport.json --output mijn_rapport.docx
        """,
    )
    parser.add_argument("json_file", help="Het JSON-bestand van de risicoanalyse")
    parser.add_argument(
        "--output", "-o",
        help="Output .docx bestand (standaard: automatisch op basis van thema)",
        default=None,
    )

    args = parser.parse_args()

    # Controleer of het bestand bestaat
    if not Path(args.json_file).exists():
        print(f"FOUT: Bestand niet gevonden: {args.json_file}")
        sys.exit(1)

    print(f"{'='*60}")
    print(f"  Cyberwijzer Rapport Generator")
    print(f"{'='*60}")
    print(f"  Input:  {args.json_file}")

    output = generate_report(args.json_file, args.output)

    print(f"  Output: {output}")
    print(f"{'='*60}")
    print(f"  Rapport succesvol gegenereerd!")


if __name__ == "__main__":
    main()
